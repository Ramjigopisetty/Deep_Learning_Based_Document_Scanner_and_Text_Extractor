from fastapi.responses import FileResponse
from docx import Document
from reportlab.pdfgen import canvas

def export_txt(text):

    file_path = "ocr_result.txt"

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text)

    return file_path


def export_docx(text):

    file_path = "ocr_result.docx"

    doc = Document()
    doc.add_heading("OCR Result", level=1)
    doc.add_paragraph(text)
    doc.save(file_path)

    return file_path


def export_pdf(text):

    file_path = "ocr_result.pdf"

    c = canvas.Canvas(file_path)
    c.drawString(100, 750, text)
    c.save()

    return file_path